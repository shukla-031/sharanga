from flask import Flask, render_template, jsonify, request, session, send_file, redirect, url_for
from sharanga_scanner import SharangaScanner
import json
import os
import hashlib
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import bcrypt
import re
import traceback
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Import database and auth
from database import *
from auth import *

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'sharanga-super-secret-key-2026')

app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
UPLOAD_FOLDER = 'uploads'
REPORTS_FOLDER = 'reports'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORTS_FOLDER, exist_ok=True)

LAST_SCAN_DATA = {}
SESSION_SCAN_DATA = {}

# ============================================================
# AUTHENTICATION ROUTES
# ============================================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        user = cursor.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
        conn.close()
        
        if user and bcrypt.checkpw(password.encode('utf-8'), user['password_hash']):
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['role'] = user['role']
            
            if user['role'] == 'admin':
                return redirect(url_for('admin_dashboard'))
            else:
                return redirect(url_for('user_dashboard'))
        else:
            return render_template('login.html', error='❌ Invalid username or password')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# ============================================================
# MAIN ROUTES
# ============================================================

@app.route('/')
def index():
    if 'user_id' in session:
        if session.get('role') == 'admin':
            return redirect(url_for('admin_dashboard'))
        else:
            return redirect(url_for('user_dashboard'))
    return redirect(url_for('login'))

@app.route('/user')
@login_required
def user_dashboard():
    return render_template('dashboard.html')

@app.route('/admin')
@admin_required
def admin_dashboard():
    return render_template('admin.html')

# ============================================================
# API ROUTES - SCAN
# ============================================================

@app.route('/api/scan', methods=['POST'])
@login_required
def scan():
    try:
        print("🔍 Scan started...")
        scanner = SharangaScanner()
        print("✅ Scanner initialized")
        
        report = scanner.scan_system()
        print("✅ Scan completed")
        
        # Get device name from report
        device_name = report.get('device', {}).get('device_name', 'Unknown')
        
        # Generate asset number using device name
        asset_number = generate_asset_number(device_name)
        report['asset_number'] = asset_number
        
        # Save to database
        save_scan_data(asset_number, report)
        print(f"✅ Data saved for asset: {asset_number}")
        
        scan_id = str(uuid.uuid4())[:8]
        report['scan_id'] = scan_id
        report['scan_time'] = datetime.now().isoformat()
        
        LAST_SCAN_DATA[scan_id] = report
        SESSION_SCAN_DATA[scan_id] = report
        
        try:
            filename = f"scan_{scan_id}.json"
            filepath = os.path.join(REPORTS_FOLDER, filename)
            with open(filepath, 'w') as f:
                json.dump(report, f, indent=2, default=str)
            
            latest_path = os.path.join(REPORTS_FOLDER, 'latest.json')
            with open(latest_path, 'w') as f:
                json.dump(report, f, indent=2, default=str)
            print(f"✅ Report saved to {filepath}")
        except Exception as e:
            print(f"⚠️ Failed to save report: {e}")
        
        return jsonify(report)
    except Exception as e:
        print(f"❌ SCAN ERROR: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/last-report')
@login_required
def last_report():
    try:
        if LAST_SCAN_DATA:
            for scan_id in sorted(LAST_SCAN_DATA.keys(), reverse=True):
                return jsonify(LAST_SCAN_DATA[scan_id])
        
        latest_path = os.path.join(REPORTS_FOLDER, 'latest.json')
        if os.path.exists(latest_path):
            with open(latest_path, 'r') as f:
                data = json.load(f)
            return jsonify(data)
    except Exception as e:
        print(f"❌ Last report error: {e}")
        print(traceback.format_exc())
    
    return jsonify({'error': 'No previous scan found'})

@app.route('/api/report/<scan_id>')
@login_required
def get_report(scan_id):
    try:
        if scan_id in LAST_SCAN_DATA:
            return jsonify(LAST_SCAN_DATA[scan_id])
        
        filepath = os.path.join(REPORTS_FOLDER, f'scan_{scan_id}.json')
        if os.path.exists(filepath):
            with open(filepath, 'r') as f:
                data = json.load(f)
            return jsonify(data)
    except Exception as e:
        print(f"❌ Report error: {e}")
        print(traceback.format_exc())
    
    return jsonify({'error': 'Scan not found'}), 404

@app.route('/api/all-reports')
@login_required
def all_reports():
    reports = []
    try:
        for filename in os.listdir(REPORTS_FOLDER):
            if filename.endswith('.json') and filename != 'latest.json':
                try:
                    with open(os.path.join(REPORTS_FOLDER, filename), 'r') as f:
                        data = json.load(f)
                        reports.append({
                            'scan_id': data.get('scan_id', filename.replace('.json', '')),
                            'scan_time': data.get('scan_time', 'Unknown'),
                            'device_name': data.get('device', {}).get('device_name', 'Unknown'),
                            'risk_score': data.get('stats', {}).get('risk_score', 0)
                        })
                except:
                    pass
    except:
        pass
    
    reports.sort(key=lambda x: x.get('scan_time', ''), reverse=True)
    return jsonify(reports)

# ============================================================
# ADMIN API ROUTES
# ============================================================

@app.route('/api/admin/assets')
@admin_required
def admin_assets():
    try:
        assets = get_all_assets()
        return jsonify(assets)
    except Exception as e:
        print(f"❌ Admin assets error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/assets/<asset_number>')
@admin_required
def admin_asset_detail(asset_number):
    try:
        asset = get_asset_by_number(asset_number)
        if asset:
            return jsonify(asset)
        return jsonify({'error': 'Asset not found'}), 404
    except Exception as e:
        print(f"❌ Asset detail error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/assets/<asset_number>/history')
@admin_required
def admin_asset_history(asset_number):
    try:
        asset = get_asset_by_number(asset_number)
        if asset:
            history = get_scan_history(asset['id'])
            return jsonify(history)
        return jsonify({'error': 'Asset not found'}), 404
    except Exception as e:
        print(f"❌ Asset history error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/users')
@admin_required
def admin_users():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        users = cursor.execute('SELECT id, username, role, email, created_at FROM users').fetchall()
        conn.close()
        return jsonify([dict(u) for u in users])
    except Exception as e:
        print(f"❌ Admin users error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/users', methods=['POST'])
@admin_required
def admin_create_user():
    try:
        data = request.json
        username = data.get('username')
        password = data.get('password')
        role = data.get('role', 'user')
        email = data.get('email', '')
        
        if not username or not password:
            return jsonify({'error': 'Username and password required'}), 400
        
        password_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO users (username, password_hash, role, email)
            VALUES (?, ?, ?, ?)
        ''', (username, password_hash, role, email))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'message': 'User created successfully'})
    except Exception as e:
        print(f"❌ Create user error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 400

@app.route('/api/admin/users/<int:user_id>', methods=['DELETE'])
@admin_required
def admin_delete_user(user_id):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM users WHERE id = ? AND role != ?', (user_id, 'admin'))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        print(f"❌ Delete user error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 400

# ============================================================
# FILE SCAN
# ============================================================

@app.route('/api/scan-file', methods=['POST'])
@login_required
def scan_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        
        file_size = os.path.getsize(filepath)
        
        md5_hash = hashlib.md5()
        sha1_hash = hashlib.sha1()
        sha256_hash = hashlib.sha256()
        
        with open(filepath, 'rb') as f:
            content = f.read()
            md5_hash.update(content)
            sha1_hash.update(content)
            sha256_hash.update(content)
        
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        suspicious_extensions = {
            '.exe': 'Executable file - HIGH RISK',
            '.scr': 'Screensaver executable - HIGH RISK',
            '.bat': 'Batch script - MEDIUM RISK',
            '.cmd': 'Command script - MEDIUM RISK',
            '.ps1': 'PowerShell script - MEDIUM RISK',
            '.vbs': 'VBScript - MEDIUM RISK',
            '.js': 'JavaScript - MEDIUM RISK',
            '.jar': 'Java executable - MEDIUM RISK',
            '.pdf': 'PDF - Can contain exploits',
            '.doc': 'Word document - Can contain macros',
            '.docx': 'Word document - Can contain macros',
            '.xls': 'Excel document - Can contain macros',
            '.xlsx': 'Excel document - Can contain macros',
            '.ppt': 'PowerPoint - Can contain macros',
            '.pptx': 'PowerPoint - Can contain macros',
            '.msi': 'Installer - MEDIUM RISK',
            '.dll': 'Dynamic Link Library - HIGH RISK',
            '.sys': 'System driver - HIGH RISK'
        }
        
        risk_level = 'LOW'
        risk_reason = 'No known risk'
        
        for ext, reason in suspicious_extensions.items():
            if file_ext == ext:
                if 'HIGH' in reason:
                    risk_level = 'HIGH'
                elif 'MEDIUM' in reason:
                    risk_level = 'MEDIUM'
                risk_reason = reason
                break
        
        file_path_lower = filepath.lower()
        suspicious_locations = ['\\temp\\', '\\downloads\\', '\\tmp\\', '\\appdata\\local\\temp\\']
        is_in_suspicious_location = any(loc in file_path_lower for loc in suspicious_locations)
        
        threat_score = 0
        if risk_level == 'HIGH':
            threat_score += 50
        elif risk_level == 'MEDIUM':
            threat_score += 25
        
        if is_in_suspicious_location:
            threat_score += 20
        
        threat_score = min(threat_score, 100)
        threat_level = 'LOW' if threat_score < 30 else 'MEDIUM' if threat_score < 60 else 'HIGH'
        
        result = {
            'filename': file.filename,
            'file_extension': file_ext,
            'file_size': f"{file_size / 1024:.2f} KB",
            'md5': md5_hash.hexdigest(),
            'sha1': sha1_hash.hexdigest(),
            'sha256': sha256_hash.hexdigest(),
            'risk_level': risk_level,
            'risk_reason': risk_reason,
            'threat_score': threat_score,
            'threat_level': threat_level,
            'suspicious_location': is_in_suspicious_location,
            'scan_time': datetime.now().isoformat(),
            'verdict': '🚨 MALICIOUS' if threat_score > 50 else '⚠️ SUSPICIOUS' if threat_score > 20 else '✅ SAFE'
        }
        
        try:
            os.remove(filepath)
        except:
            pass
        
        return jsonify(result)
    except Exception as e:
        print(f"❌ File scan error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/clear')
@login_required
def clear_data():
    global LAST_SCAN_DATA, SESSION_SCAN_DATA
    LAST_SCAN_DATA = {}
    SESSION_SCAN_DATA = {}
    return jsonify({'status': 'cleared'})

# ============================================================
# DOWNLOAD INDIVIDUAL USER REPORT
# ============================================================

@app.route('/api/download-individual-report', methods=['POST'])
@admin_required
def download_individual_report():
    try:
        data = request.json
        asset_number = data.get('asset_number')
        
        if not asset_number:
            return jsonify({'error': 'Asset number required'}), 400
        
        asset = get_asset_by_number(asset_number)
        if not asset:
            return jsonify({'error': 'Asset not found'}), 404
        
        scan_data = json.loads(asset['scan_data']) if asset['scan_data'] else {}
        device = scan_data.get('device', {})
        config = scan_data.get('system_config', {})
        network = scan_data.get('network', {})
        apps = scan_data.get('installed_apps', [])
        processes = scan_data.get('running_processes', [])
        stats = scan_data.get('stats', {})
        
        doc = Document()
        
        title = doc.add_heading('🏹 Sharanga - Individual Security Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        sub = doc.add_paragraph(f'Asset: {asset_number}')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('-' * 50)
        
        # Device Information
        doc.add_heading('📱 Device Information', level=1)
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info = [
            ('Device Name', device.get('device_name', 'Unknown')),
            ('User Name', device.get('username', 'Unknown')),
            ('Domain', device.get('domain', 'Unknown')),
            ('Domain Joined', 'Yes' if device.get('is_domain_joined') else 'No'),
            ('Operating System', f"{device.get('os_name', 'Unknown')} {device.get('os_release', '')}"),
            ('IP Address', device.get('ip_address', 'Unknown')),
            ('RAM', f"{device.get('ram_gb', 0)} GB"),
            ('System Manufacturer', device.get('system_manufacturer', 'Unknown')),
            ('System Model', device.get('system_model', 'Unknown')),
            ('Uptime', config.get('uptime', 'Unknown'))
        ]
        
        for i, (label, value) in enumerate(info):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # CPU
        cpu = config.get('cpu', {})
        if cpu:
            doc.add_heading('⚙️ CPU Information', level=1)
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light Grid Accent 1'
            
            cpu_info = [
                ('CPU Name', cpu.get('name', 'Unknown')),
                ('Physical Cores', str(cpu.get('cores', 'Unknown'))),
                ('Logical Processors', str(cpu.get('logical', 'Unknown'))),
                ('Max Clock Speed', cpu.get('max_clock', 'Unknown')),
                ('Current Clock Speed', cpu.get('current_clock', 'Unknown')),
                ('Current Usage', f"{cpu.get('usage', 0)}%")
            ]
            
            for i, (label, value) in enumerate(cpu_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Memory
        memory = config.get('memory', {})
        if memory:
            doc.add_heading('💾 Memory Information', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'
            
            mem_info = [
                ('Total RAM', f"{memory.get('total_gb', 0)} GB"),
                ('Used RAM', f"{memory.get('used_gb', 0)} GB"),
                ('Available RAM', f"{memory.get('available_gb', 0)} GB"),
                ('Usage', f"{memory.get('used_percent', 0)}%")
            ]
            
            for i, (label, value) in enumerate(mem_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Motherboard
        motherboard = config.get('motherboard', {})
        if motherboard:
            doc.add_heading('🖥️ Motherboard Information', level=1)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'
            
            mb_info = [
                ('Motherboard Name', motherboard.get('name', 'Unknown')),
                ('Manufacturer', motherboard.get('manufacturer', 'Unknown'))
            ]
            
            for i, (label, value) in enumerate(mb_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # BIOS
        bios = config.get('bios', {})
        if bios:
            doc.add_heading('🔧 BIOS Information', level=1)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'
            
            bios_info = [
                ('BIOS Version', bios.get('version', 'Unknown')),
                ('Manufacturer', bios.get('manufacturer', 'Unknown'))
            ]
            
            for i, (label, value) in enumerate(bios_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Graphics
        gpus = config.get('gpus', [])
        if gpus:
            doc.add_heading('🎮 Graphics Information', level=1)
            table = doc.add_table(rows=1+len(gpus), cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = ['GPU Name', 'VRAM', 'Driver']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for idx, gpu in enumerate(gpus):
                table.rows[idx+1].cells[0].text = gpu.get('name', 'Unknown')
                table.rows[idx+1].cells[1].text = gpu.get('ram', 'Unknown')
                table.rows[idx+1].cells[2].text = gpu.get('driver', 'Unknown')
            
            doc.add_paragraph()
        
        # Network
        if network:
            doc.add_heading('🌐 Network Information', level=1)
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Light Grid Accent 1'
            
            net_info = [
                ('Connection Type', network.get('connection_type', 'Unknown')),
                ('Interface', network.get('interface', 'Unknown')),
                ('MAC Address', network.get('mac_address', 'Unknown')),
                ('IP Address', network.get('ip_address', 'Unknown')),
                ('SSID', network.get('ssid', 'Unknown')),
                ('Speed', network.get('speed', 'Unknown')),
                ('Gateway', network.get('gateway', 'Unknown'))
            ]
            
            for i, (label, value) in enumerate(net_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Weather
        weather = config.get('weather', {})
        if weather:
            doc.add_heading('🌤️ Weather Information', level=1)
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'
            
            weather_info = [
                ('Location', weather.get('city', 'Unknown')),
                ('Temperature', weather.get('temperature', '--°C')),
                ('Condition', weather.get('condition', 'Unknown')),
                ('Humidity', weather.get('humidity', '--%')),
                ('Wind Speed', weather.get('wind_speed', '-- km/h'))
            ]
            
            for i, (label, value) in enumerate(weather_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Disk Drives
        disks = config.get('disks', [])
        if disks:
            doc.add_heading('💾 Disk Drives', level=1)
            table = doc.add_table(rows=1+len(disks), cols=5)
            table.style = 'Light Grid Accent 1'
            
            headers = ['Drive', 'Total (GB)', 'Used (GB)', 'Free (GB)', 'Usage %']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for row_idx, disk in enumerate(disks):
                table.rows[row_idx+1].cells[0].text = disk.get('drive', 'Unknown')
                table.rows[row_idx+1].cells[1].text = str(disk.get('total_gb', 0))
                table.rows[row_idx+1].cells[2].text = str(disk.get('used_gb', 0))
                table.rows[row_idx+1].cells[3].text = str(disk.get('free_gb', 0))
                table.rows[row_idx+1].cells[4].text = f"{disk.get('used_percent', 0)}%"
            
            doc.add_paragraph()
        
        # ========== INSTALLED APPLICATIONS (ALL - NO LIMIT) ==========
        doc.add_heading(f'📦 Installed Applications ({len(apps)})', level=1)
        
        if apps:
            table = doc.add_table(rows=1+len(apps), cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = ['#', 'Application Name', 'Version']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for idx, app in enumerate(apps):
                table.rows[idx+1].cells[0].text = str(idx+1)
                table.rows[idx+1].cells[1].text = app.get('name', 'Unknown')
                table.rows[idx+1].cells[2].text = app.get('version', 'Unknown')
        else:
            doc.add_paragraph('No applications found')
        
        doc.add_paragraph()
        
        # Suspicious Processes
        suspicious = [p for p in processes if p.get('is_suspicious')]
        if suspicious:
            doc.add_heading(f'⚠️ Suspicious Processes ({len(suspicious)})', level=1)
            table = doc.add_table(rows=1+len(suspicious), cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = ['Process Name', 'PID', 'Status']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for idx, p in enumerate(suspicious):
                table.rows[idx+1].cells[0].text = p.get('name', 'Unknown')
                table.rows[idx+1].cells[1].text = str(p.get('pid', 'Unknown'))
                table.rows[idx+1].cells[2].text = '⚠️ Suspicious'
            
            doc.add_paragraph()
        
        # Battery
        if config.get('battery_percent'):
            doc.add_paragraph(f'🔋 Battery: {config.get("battery_percent")}% ({config.get("battery_status", "Unknown")})')
            doc.add_paragraph()
        
        # Summary
        doc.add_heading('📊 Summary', level=1)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        risk_score = stats.get('risk_score', 0)
        risk_level = 'HIGH' if risk_score >= 60 else 'MEDIUM' if risk_score >= 30 else 'LOW'
        
        summary_data = [
            ('Total Applications', str(len(apps))),
            ('Running Processes', str(len(processes))),
            ('Suspicious Processes', str(len(suspicious))),
            ('Risk Score', f"{risk_score}/100"),
            ('Risk Level', risk_level)
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Footer
        doc.add_paragraph('-' * 50)
        footer = doc.add_paragraph('🏹 Sharanga - The Divine Bow of Cyber Defense')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].italic = True
        footer = doc.add_paragraph('Generated by Sharanga Security Tool')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"sharanga_{asset_number}_{timestamp}.docx"
        
        file_buffer = io.BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)
        
        return send_file(
            file_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"❌ Individual report error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

# ============================================================
# DOWNLOAD FULL REPORT
# ============================================================

@app.route('/api/download-report')
@admin_required
def download_report():
    try:
        latest_path = os.path.join(REPORTS_FOLDER, 'latest.json')
        if not os.path.exists(latest_path):
            return jsonify({'error': 'No report found'}), 404
        
        with open(latest_path, 'r') as f:
            data = json.load(f)
        
        doc = Document()
        
        title = doc.add_heading('🏹 Sharanga - Security Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        sub = doc.add_paragraph('The Divine Bow of Cyber Defense')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True
        
        doc.add_paragraph(f'Scan Completed: {data.get("scan_time", "Unknown")}')
        doc.add_paragraph('-' * 50)
        
        # Device Information
        doc.add_heading('📱 Device Information', level=1)
        device = data.get('device', {})
        table = doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info = [
            ('Device Name', device.get('device_name', 'Unknown')),
            ('User Name', device.get('username', 'Unknown')),
            ('Domain', device.get('domain', 'Unknown')),
            ('Domain Joined', 'Yes' if device.get('is_domain_joined') else 'No'),
            ('Operating System', f"{device.get('os_name', 'Unknown')} {device.get('os_release', '')}"),
            ('IP Address', device.get('ip_address', 'Unknown')),
            ('RAM', f"{device.get('ram_gb', 0)} GB"),
            ('Processor', device.get('processor', 'Unknown')),
            ('System Manufacturer', device.get('system_manufacturer', 'Unknown')),
            ('System Model', device.get('system_model', 'Unknown')),
            ('Uptime', data.get('system_config', {}).get('uptime', 'Unknown'))
        ]
        
        for i, (label, value) in enumerate(info):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # CPU
        config = data.get('system_config', {})
        cpu = config.get('cpu', {})
        if cpu:
            doc.add_heading('⚙️ CPU Information', level=1)
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light Grid Accent 1'
            
            cpu_info = [
                ('CPU Name', cpu.get('name', 'Unknown')),
                ('Physical Cores', str(cpu.get('cores', 'Unknown'))),
                ('Logical Processors', str(cpu.get('logical', 'Unknown'))),
                ('Max Clock Speed', cpu.get('max_clock', 'Unknown')),
                ('Current Clock Speed', cpu.get('current_clock', 'Unknown')),
                ('Current Usage', f"{cpu.get('usage', 0)}%")
            ]
            
            for i, (label, value) in enumerate(cpu_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Memory
        memory = config.get('memory', {})
        if memory:
            doc.add_heading('💾 Memory Information', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'
            
            mem_info = [
                ('Total RAM', f"{memory.get('total_gb', 0)} GB"),
                ('Used RAM', f"{memory.get('used_gb', 0)} GB"),
                ('Available RAM', f"{memory.get('available_gb', 0)} GB"),
                ('Usage', f"{memory.get('used_percent', 0)}%")
            ]
            
            for i, (label, value) in enumerate(mem_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Motherboard
        motherboard = config.get('motherboard', {})
        if motherboard:
            doc.add_heading('🖥️ Motherboard Information', level=1)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'
            
            mb_info = [
                ('Motherboard Name', motherboard.get('name', 'Unknown')),
                ('Manufacturer', motherboard.get('manufacturer', 'Unknown'))
            ]
            
            for i, (label, value) in enumerate(mb_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # BIOS
        bios = config.get('bios', {})
        if bios:
            doc.add_heading('🔧 BIOS Information', level=1)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'
            
            bios_info = [
                ('BIOS Version', bios.get('version', 'Unknown')),
                ('Manufacturer', bios.get('manufacturer', 'Unknown'))
            ]
            
            for i, (label, value) in enumerate(bios_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Graphics
        gpus = config.get('gpus', [])
        if gpus:
            doc.add_heading('🎮 Graphics Information', level=1)
            table = doc.add_table(rows=1+len(gpus), cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = ['GPU Name', 'VRAM', 'Driver']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for idx, gpu in enumerate(gpus):
                table.rows[idx+1].cells[0].text = gpu.get('name', 'Unknown')
                table.rows[idx+1].cells[1].text = gpu.get('ram', 'Unknown')
                table.rows[idx+1].cells[2].text = gpu.get('driver', 'Unknown')
            
            doc.add_paragraph()
        
        # Network
        network = data.get('network', {})
        if network:
            doc.add_heading('🌐 Network Information', level=1)
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Light Grid Accent 1'
            
            net_info = [
                ('Connection Type', network.get('connection_type', 'Unknown')),
                ('Interface', network.get('interface', 'Unknown')),
                ('MAC Address', network.get('mac_address', 'Unknown')),
                ('IP Address', network.get('ip_address', 'Unknown')),
                ('SSID', network.get('ssid', 'Unknown')),
                ('Speed', network.get('speed', 'Unknown')),
                ('Gateway', network.get('gateway', 'Unknown'))
            ]
            
            for i, (label, value) in enumerate(net_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Weather
        weather = config.get('weather', {})
        if weather:
            doc.add_heading('🌤️ Weather Information', level=1)
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light Grid Accent 1'
            
            weather_info = [
                ('Location', weather.get('city', 'Unknown')),
                ('Temperature', weather.get('temperature', '--°C')),
                ('Condition', weather.get('condition', 'Unknown')),
                ('Humidity', weather.get('humidity', '--%')),
                ('Wind Speed', weather.get('wind_speed', '-- km/h')),
                ('Pressure', weather.get('pressure', '-- hPa'))
            ]
            
            for i, (label, value) in enumerate(weather_info):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            doc.add_paragraph()
        
        # Disk Drives
        disks = config.get('disks', [])
        if disks:
            doc.add_heading('💾 Disk Drives', level=1)
            table = doc.add_table(rows=1+len(disks), cols=5)
            table.style = 'Light Grid Accent 1'
            
            headers = ['Drive', 'Total (GB)', 'Used (GB)', 'Free (GB)', 'Usage %']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for row_idx, disk in enumerate(disks):
                table.rows[row_idx+1].cells[0].text = disk.get('drive', 'Unknown')
                table.rows[row_idx+1].cells[1].text = str(disk.get('total_gb', 0))
                table.rows[row_idx+1].cells[2].text = str(disk.get('used_gb', 0))
                table.rows[row_idx+1].cells[3].text = str(disk.get('free_gb', 0))
                table.rows[row_idx+1].cells[4].text = f"{disk.get('used_percent', 0)}%"
            
            doc.add_paragraph()
        
        # Battery
        if config.get('battery_percent'):
            doc.add_paragraph(f'🔋 Battery: {config.get("battery_percent")}% ({config.get("battery_status", "Unknown")})')
            doc.add_paragraph()
        
        # ========== INSTALLED APPLICATIONS (ALL - NO LIMIT) ==========
        apps = data.get('installed_apps', [])
        doc.add_heading(f'📦 Installed Applications ({len(apps)})', level=1)
        
        if apps:
            table = doc.add_table(rows=1+len(apps), cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = ['#', 'Application Name', 'Version']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for idx, app in enumerate(apps):
                table.rows[idx+1].cells[0].text = str(idx+1)
                table.rows[idx+1].cells[1].text = app.get('name', 'Unknown')
                table.rows[idx+1].cells[2].text = app.get('version', 'Unknown')
        else:
            doc.add_paragraph('No applications found')
        
        doc.add_paragraph()
        
        # Suspicious Processes
        processes = data.get('running_processes', [])
        suspicious = [p for p in processes if p.get('is_suspicious')]
        if suspicious:
            doc.add_heading(f'⚠️ Suspicious Processes ({len(suspicious)})', level=1)
            table = doc.add_table(rows=1+len(suspicious), cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = ['Process Name', 'PID', 'Status']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for idx, p in enumerate(suspicious):
                table.rows[idx+1].cells[0].text = p.get('name', 'Unknown')
                table.rows[idx+1].cells[1].text = str(p.get('pid', 'Unknown'))
                table.rows[idx+1].cells[2].text = '⚠️ Suspicious'
            
            doc.add_paragraph()
        
        # Uninstalled
        uninstalled = data.get('uninstalled_apps', [])
        if uninstalled:
            doc.add_heading(f'🗑️ Recently Uninstalled ({len(uninstalled)})', level=1)
            table = doc.add_table(rows=1+len(uninstalled), cols=2)
            table.style = 'Light Grid Accent 1'
            
            headers = ['Application Name', 'Uninstalled Date']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            
            for idx, app in enumerate(uninstalled):
                table.rows[idx+1].cells[0].text = app.get('app_name', 'Unknown')
                table.rows[idx+1].cells[1].text = str(app.get('uninstall_date', 'Unknown'))
            
            doc.add_paragraph()
        
        # Summary
        stats = data.get('stats', {})
        doc.add_heading('📊 Summary', level=1)
        
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        risk_score = stats.get('risk_score', 0)
        risk_level = 'HIGH' if risk_score >= 60 else 'MEDIUM' if risk_score >= 30 else 'LOW'
        
        summary_data = [
            ('Total Applications', str(len(apps))),
            ('Running Processes', str(len(processes))),
            ('Suspicious Processes', str(len(suspicious))),
            ('Uninstalled (7 days)', str(stats.get('total_uninstalled', 0))),
            ('Risk Score', f"{risk_score}/100"),
            ('Risk Level', risk_level)
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Footer
        doc.add_paragraph('-' * 50)
        footer = doc.add_paragraph('🏹 Sharanga - The Divine Bow of Cyber Defense')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].italic = True
        footer = doc.add_paragraph('Generated by Sharanga Security Tool')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"sharanga_report_{timestamp}.docx"
        
        file_buffer = io.BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)
        
        return send_file(
            file_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"❌ Download report error: {e}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

# ============================================================
# RUN APPLICATION (FIXED FOR RENDER)
# ============================================================

if __name__ == '__main__':
    init_db()
    # Render uses PORT environment variable
    port = int(os.environ.get('PORT', 5000))
    # Debug mode only on local, not on Render
    debug = os.environ.get('FLASK_ENV') != 'production'
    app.run(host='0.0.0.0', port=port, debug=debug)